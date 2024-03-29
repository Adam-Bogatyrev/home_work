from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
doc = Document()
doc.add_paragraph("Hello Python")
doc.save("hello_python.docx")
loaded_doc = Document("hello_python.docx")
bold_text = ""
for paragraph in loaded_doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text
print("Bold Text:", bold_text)
new_doc = Document()
paragraph = new_doc.add_paragraph("New paragraph with different font.")
run = paragraph.run